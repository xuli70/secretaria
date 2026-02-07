import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_SYSTEM_PROMPT = (
    "Eres Secretaria, un asistente especializado en redactar documentos formales. "
    "Genera el documento solicitado con formato profesional. "
    "Usa lineas que empiecen con # para titulos, ## para subtitulos, "
    "- o * para listas, y texto normal para parrafos. "
    "No incluyas explicaciones fuera del documento, solo el contenido del documento."
)


def generate_docx(content: str, title: str, save_dir: str) -> tuple[str, str]:
    """Convert text content to a DOCX file.

    Returns (filepath, filename).
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Parse content line by line
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()

        if not stripped:
            # Empty line → small paragraph break
            doc.add_paragraph("")
            continue

        # Heading level 2: ## text
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_heading(text, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # Heading level 1: # text
        if stripped.startswith("# "):
            text = stripped[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # List item: - text or * text
        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            doc.add_paragraph(text, style="List Bullet")
            continue

        # Normal paragraph
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Generate filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"doc_{timestamp}.docx"
    filepath = os.path.join(save_dir, filename)

    os.makedirs(save_dir, exist_ok=True)
    doc.save(filepath)

    return filepath, filename


def generate_txt(content: str, title: str, save_dir: str) -> tuple[str, str]:
    """Save content as a plain text file.

    Returns (filepath, filename).
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"doc_{timestamp}.txt"
    filepath = os.path.join(save_dir, filename)

    os.makedirs(save_dir, exist_ok=True)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    return filepath, filename
